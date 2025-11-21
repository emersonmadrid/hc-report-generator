from pathlib import Path
from typing import List, Dict

from docx import Document


def generar_reporte(
    bloques: List[Dict[str, List[str]]],
    plantilla_path: Path,
    output_path: Path,
) -> None:
    """
    Crea un Word a partir de una plantilla y una lista de bloques de datos.
    Cada bloque corresponde a un Excel.
    Agrupa los datos en tablas de máximo 6 columnas de datos.
    """
    doc = Document(plantilla_path)

    # Consolidar todos los datos de todos los archivos Excel
    todos_hc = []
    todas_fechas = []
    todos_porcentajes = []
    todas_calificaciones = []

    for bloque in bloques:
        todos_hc.extend(bloque["hc"])
        todas_fechas.extend(bloque["fechas"])
        todos_porcentajes.extend(bloque["porcentajes"])
        todas_calificaciones.extend(bloque["calificaciones"])

    print(f"\n=== Generando reporte ===")
    print(f"Total de registros: {len(todos_hc)}")

    # Agrupar los datos en bloques de 6
    MAX_COLS_POR_TABLA = 6
    tabla_idx = 0

    for i in range(0, len(todos_hc), MAX_COLS_POR_TABLA):
        # Tomar hasta 6 registros
        hc_grupo = todos_hc[i:i + MAX_COLS_POR_TABLA]
        fechas_grupo = todas_fechas[i:i + MAX_COLS_POR_TABLA]
        porc_grupo = todos_porcentajes[i:i + MAX_COLS_POR_TABLA]
        calif_grupo = todas_calificaciones[i:i + MAX_COLS_POR_TABLA]

        print(f"  Tabla {tabla_idx + 1}: {len(hc_grupo)} registros")

        # Si hay tablas existentes en la plantilla, usarlas
        if tabla_idx < len(doc.tables):
            tabla = doc.tables[tabla_idx]
            print(f"    → Usando tabla existente {tabla_idx + 1}")
        else:
            # Si ya no hay tablas en la plantilla, crear una nueva
            print(f"    → Creando nueva tabla")
            # Agregar un título antes de la tabla
            doc.add_paragraph("EVALUACIÓN DE LA CALIDAD DE REGISTRO EN CONSULTA EXTERNA")
            
            num_cols = len(hc_grupo) + 1  # +1 para columna de títulos
            tabla = doc.add_table(rows=4, cols=num_cols)
            tabla.style = 'Table Grid'

        # Asegurar que la tabla tenga al menos 4 filas
        while len(tabla.rows) < 4:
            tabla.add_row()

        # Calcular cuántas columnas necesitamos
        num_cols_necesarias = len(hc_grupo) + 1

        # Agregar columnas si hacen falta
        while len(tabla.columns) < num_cols_necesarias:
            tabla.add_column(width=1)

        # Limpiar todas las celdas primero
        for row_idx in range(4):
            for col_idx in range(len(tabla.columns)):
                tabla.rows[row_idx].cells[col_idx].text = ""

        # Llenar la primera columna con los títulos
        tabla.rows[0].cells[0].text = "N° de HC evaluada"
        tabla.rows[1].cells[0].text = "Fecha de atención evaluada"
        tabla.rows[2].cells[0].text = "% cumplimiento"
        tabla.rows[3].cells[0].text = "Calificación del registro"

        # Llenar las columnas de datos
        for idx, (h, f, p, c) in enumerate(
            zip(hc_grupo, fechas_grupo, porc_grupo, calif_grupo), start=1
        ):
            if idx >= len(tabla.columns):
                tabla.add_column(width=1)

            tabla.rows[0].cells[idx].text = h
            tabla.rows[1].cells[idx].text = f
            tabla.rows[2].cells[idx].text = p
            tabla.rows[3].cells[idx].text = c

        tabla_idx += 1

        # Agregar un salto de línea entre tablas si no es la última
        if i + MAX_COLS_POR_TABLA < len(todos_hc):
            doc.add_paragraph("")

    print(f"  Total de tablas creadas/llenadas: {tabla_idx}")

    doc.save(output_path)
    print(f"✓ Reporte guardado en: {output_path}")